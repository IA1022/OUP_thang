from collections import defaultdict
import re
from docx import Document
from lxml import etree

# QUESTION_START updated: require some text after the number if it's a numbered question,
# but we'll also disambiguate numbered option lines in the parser using context.
QUESTION_START = re.compile(r"(?i)^(question(\s*\d+)?)\b|^(\d+[\.\)]\s+\S.*)")

# NON-MCQ detection (case-insensitive, fuzzy)
NON_MCQ_SECTION = re.compile(
    r"(?i)(essay|true\s*or\s*false|true/false|short\s*answer|long\s*answer)"
)

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _get_numbering_part_xml(doc):
    for part in doc.part.package.parts:
        if (
            part.content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"
        ):
            return etree.fromstring(part.blob)
    return None


def get_level_format_from_numId(doc, numId, ilvl):
    numId = str(numId)
    numbering_xml = _get_numbering_part_xml(doc)
    if numbering_xml is None:
        return None, None

    num_elm = numbering_xml.find(f".//w:num[@w:numId='{numId}']", namespaces=W_NS)
    if num_elm is None:
        return None, None

    abs_node = num_elm.find("./w:abstractNumId", namespaces=W_NS)
    if abs_node is None:
        return None, None
    abstract_num_id = abs_node.get(f"{{{W_NS['w']}}}val")

    abs_elm = numbering_xml.find(
        f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']", namespaces=W_NS
    )
    if abs_elm is None:
        return None, None

    lvl_elm = abs_elm.find(f"./w:lvl[@w:ilvl='{ilvl}']", namespaces=W_NS)
    if lvl_elm is None:
        return None, None

    numFmt_elm = lvl_elm.find("./w:numFmt", namespaces=W_NS)
    lvlText_elm = lvl_elm.find("./w:lvlText", namespaces=W_NS)

    numFmt = numFmt_elm.get(f"{{{W_NS['w']}}}val") if numFmt_elm is not None else None
    lvlText = (
        lvlText_elm.get(f"{{{W_NS['w']}}}val") if lvlText_elm is not None else None
    )

    return numFmt, lvlText


def convert_level_to_number(n, fmt):
    if fmt == "decimal":
        return str(n)
    if fmt == "lowerLetter":
        return chr(ord("a") + (n - 1))
    if fmt == "upperLetter":
        return chr(ord("A") + (n - 1))
    if fmt == "lowerRoman":
        return to_roman(n).lower()
    if fmt == "upperRoman":
        return to_roman(n)
    return str(n)


def to_roman(n):
    vals = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    res = ""
    for v, r in vals:
        while n >= v:
            res += r
            n -= v
    return res


# Helper regexes
OPTION_NUMERIC = re.compile(r"^\s*\d+[\.\)]\s+(.+)$")  # "1. text" or "2) text"
OPTION_LETTER = re.compile(r"^\s*[a-eA-E][\.\)]\s*(.+)$")  # "a. text" or "B) text"
PLAIN_LINE = re.compile(r"^\s*(.+\S)\s*$")  # non-empty line


def clean_option_text(t):
    """Strip option prefixes (1., a.) and return clean option text."""
    if t is None:
        return ""
    t = t.strip()
    m = OPTION_LETTER.match(t)
    if m:
        return m.group(1).strip()
    m = OPTION_NUMERIC.match(t)
    if m:
        return m.group(1).strip()
    return t


def parse_answer_key(lines):
    answers = []
    start = None

    for i, ln in enumerate(lines):
        if re.search(r"(?i)^answer\s*key", ln):
            start = i + 1
            break

    if start is None:
        return answers

    for ln in lines[start:]:
        ln = ln.strip()
        if not ln:
            continue

        # Stop if non-MCQ section appears
        if NON_MCQ_SECTION.search(ln):
            break

        # Single-letter answers
        if re.fullmatch(r"[A-Ea-e]", ln):
            answers.append(ln.lower())
            continue

        # formats like "1. a" or "2) B"
        m = re.match(r"\d+[\.\)]\s*([a-eA-E])", ln)
        if m:
            answers.append(m.group(1).lower())
            continue

        # Sometimes answer list may be like "1 a" (space)
        m = re.match(r"\d+\s+([a-eA-E])", ln)
        if m:
            answers.append(m.group(1).lower())
            continue

    return answers


def parse_questions(lines):
    """
    Parse MCQs from the lines list.
    Stop entirely when a NON_MCQ_SECTION line is encountered.
    Treat numbered lines (1.,2.) as options if we currently have an open question.
    """
    qs = []
    current = None
    qnum = 1

    i = 0
    while i < len(lines):
        ln = lines[i]

        # If we reach essay/true-false, stop parsing questions completely
        if NON_MCQ_SECTION.search(ln) or re.search(r"(?i)^answer\s*key", ln):
            break

        # If line starts a question (QUESTION_START) and we don't already treat it as an option
        qstart_match = QUESTION_START.search(ln)
        opt_num_match = OPTION_NUMERIC.match(ln)
        opt_letter_match = OPTION_LETTER.match(ln)

        # If we currently have an open question:
        if current:
            # If this line is a numeric option or letter option → append as option
            if opt_num_match:
                current["options"].append(clean_option_text(ln))
                i += 1
                continue
            if opt_letter_match:
                current["options"].append(clean_option_text(ln))
                i += 1
                continue

            # If line looks like a new question start (and not an option) → close current
            if qstart_match:
                qs.append(current)
                current = None
                # do not advance here, let the next iteration start this new question
                continue

            # Otherwise it's a continuation of question text or continuation of last option
            if current["options"]:
                # continuation of last option
                current["options"][-1] = (
                    current["options"][-1] + " " + ln.strip()
                ).strip()
            else:
                # continuation of question text
                current["question"] = (current["question"] + " " + ln.strip()).strip()
            i += 1
            continue

        # If no current question:
        if qstart_match:
            # But if the match is a numeric line and actually is an option-ish (rare here),
            # we only start a question if the line has significant text (QUESTION_START regex ensures that).
            current = {"number": qnum, "question": ln.strip(), "options": []}
            qnum += 1
            i += 1
            continue

        # If this line matches an option but we have no current question, it's ambiguous:
        # treat it as part of previous question only if there is a previous question in qs.
        if opt_num_match or opt_letter_match:
            if qs:
                # append to last question's options
                last = qs[-1]
                last_opts = last.get("options", [])
                last_opts.append(clean_option_text(ln))
                last["options"] = last_opts
            else:
                # no question to attach to — ignore or create a dummy question? we'll create a small placeholder
                current = {
                    "number": qnum,
                    "question": "",
                    "options": [clean_option_text(ln)],
                }
                qnum += 1
            i += 1
            continue

        # Otherwise skip stray lines
        i += 1

    # append last open question
    if current:
        qs.append(current)

    # Normalize options list to always have 5 slots (a-e)
    for q in qs:
        opts = q.get("options", [])
        # remove empty options that are empty strings
        opts = [o for o in opts]
        while len(opts) < 5:
            opts.append("")
        q["options"] = opts[:5]

    return qs


def create_output_doc(
    questions, answers, non_mcq_start_index, original_lines, out_path
):
    doc = Document()

    # Write MCQs
    for q in questions:
        qnum = q["number"]
        qtext = q["question"]
        doc.add_paragraph(f"Question {qnum}) {qtext}")
        letters = ["a", "b", "c", "d", "e"]
        for letter, opt in zip(letters, q["options"]):
            doc.add_paragraph(f"{letter}. {opt}")
        if qnum - 1 < len(answers):
            doc.add_paragraph(f"Answer: {answers[qnum - 1]}")
        doc.add_paragraph("")

    # Append non-MCQ content verbatim (essay / true-false / rest of document)
    if non_mcq_start_index is not None and non_mcq_start_index < len(original_lines):
        doc.add_page_break()
        for ln in original_lines[non_mcq_start_index:]:
            doc.add_paragraph(ln)

    doc.save(out_path)


def process_docx(path, output_path):
    doc = Document(path)
    lines = []
    non_mcq_index = None

    # Build `lines` with numbering resolved; remember first NON_MCQ paragraph index
    for p in doc.paragraphs:
        text = p.text or ""
        # If we've already flagged start of non-mcq, just keep collecting lines
        if non_mcq_index is not None:
            lines.append(text)
            continue

        # If this paragraph indicates non-mcq section, mark the index and append
        if NON_MCQ_SECTION.search(text):
            non_mcq_index = len(lines)
            lines.append(text)
            continue

        # Always append paragraph text; numbering logic below tries to resolve numbered lists into e.g. "1. text"
        p_elm = p._p
        numPr = None
        if p_elm.pPr is not None and p_elm.pPr.numPr is not None:
            numPr = p_elm.pPr.numPr

        if numPr is None:
            lines.append(text)
            continue

        # If numbered paragraph: resolve prefix using numbering part (as you did earlier)
        numId = numPr.numId.val
        ilvl = int(numPr.ilvl.val) if numPr.ilvl is not None else 0

        # counters per numId/ilvl to build numeric indexes
        # reuse counters like your earlier script
        # we'll maintain counters dict in function-local scope
        # (initialize when first used)
        if "counters" not in locals():
            counters = defaultdict(lambda: defaultdict(int))
            last_level_for_num = defaultdict(lambda: -1)
        if last_level_for_num[numId] == -1 or ilvl <= last_level_for_num[numId]:
            for deeper in list(counters[numId].keys()):
                if deeper > ilvl:
                    del counters[numId][deeper]
        counters[numId][ilvl] += 1
        last_level_for_num[numId] = ilvl

        numFmt, lvlText = get_level_format_from_numId(doc, numId, ilvl)

        if numFmt == "bullet" or (
            lvlText is not None and "%" not in (lvlText or "") and numFmt == "bullet"
        ):
            prefix = (lvlText or "•") + " "
        else:
            val = convert_level_to_number(counters[numId][ilvl], numFmt)
            if lvlText:
                placeholder = f"%{ilvl + 1}"
                prefix = (
                    (lvlText.replace(placeholder, val) + " ")
                    if "%" in (lvlText or "")
                    else (val + ". ")
                )
            else:
                prefix = val + ". "

        lines.append(prefix + text)

    # Ensure non_mcq_index set (if no non-mcq section found, it will be len(lines))
    if non_mcq_index is None:
        non_mcq_index = len(lines)

    # Parse answer key from full lines (answer key always at end)
    answers = parse_answer_key(lines)

    # Parse questions only up to non_mcq_index (stop before essay/true-false)
    questions = parse_questions(lines[:non_mcq_index])

    create_output_doc(questions, answers, non_mcq_index, lines, output_path)
    print("Document processed and saved.")


if __name__ == "__main__":
    input_file = "C:\\Users\\nikhi\\OneDrive\\Desktop\\Scripts\\test_new.docx"
    output_file = (
        "C:\\Users\\nikhi\\OneDrive\\Desktop\\Scripts\\test_new_processed13.docx"
    )
    process_docx(input_file, output_file)
